# -*- coding: utf-8 -*-
"""
Genera el documento Paper de la simulación según formato UTN (formato_papers_estudiantes).
Usa los datos del benchmark corrido. Formato: 2 columnas, espaciado 1 cm.
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

# Configuración de estilos según formato UTN
TNR = "Times New Roman"
TITULO_SIZE = Pt(16)
AUTOR_SIZE = Pt(14)
SECCION_SIZE = Pt(12)
TEXTO_SIZE = Pt(12)
ABSTRACT_SIZE = Pt(10)
REF_SIZE = Pt(10)


def _agregar_section_break_2_columnas(doc):
    """Inserta un section break y configura la siguiente sección con 2 columnas, 1 cm entre columnas."""
    # 1 cm = 567 twips (1 inch = 1440 twips)
    p = doc.add_paragraph()
    sectPr = OxmlElement("w:sectPr")
    cols = OxmlElement("w:cols")
    cols.set(qn("w:num"), "2")
    cols.set(qn("w:space"), "567")  # 1 cm entre columnas
    sectPr.append(cols)
    p._p.append(sectPr)


def crear_paper():
    doc = Document()
    
    # Configurar márgenes A4: 2.5 cm
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
    
    # --- TÍTULO (Times New Roman, 16, negrita) ---
    # Alineado con propuesta: "Análisis del comportamiento de una plataforma SaaS de soporte técnico bajo demanda"
    titulo = doc.add_paragraph()
    titulo_run = titulo.add_run(
        "Análisis del Comportamiento de una Plataforma SaaS de Soporte Técnico Bajo Demanda "
        "a Través de la Simulación: Benchmark de Viabilidad Económica a 10 Años"
    )
    titulo_run.bold = True
    titulo_run.font.size = TITULO_SIZE
    titulo_run.font.name = TNR
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # --- AUTORES (Times New Roman, 14, negrita) con legajos ---
    autores = doc.add_paragraph()
    autores_run = autores.add_run(
        "Luciano Sonntag, Claudio Hernán (120580-8); Pacay, Ariel Gastón (121237-0); Reynoso, Alison Michella (156899-1)"
    )
    autores_run.bold = True
    autores_run.font.size = AUTOR_SIZE
    autores_run.font.name = TNR
    autores.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # --- INSTITUCIÓN y CÁTEDRA ---
    inst = doc.add_paragraph()
    inst_run = inst.add_run("Universidad Tecnológica Nacional, Facultad Regional Buenos Aires")
    inst_run.bold = True
    inst_run.italic = True
    inst_run.font.size = Pt(12)
    inst_run.font.name = TNR
    inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    catedra = doc.add_paragraph()
    cat_run = catedra.add_run("Ingeniería en Sistemas de Información – Cátedra: Simulación")
    cat_run.bold = True
    cat_run.italic = True
    cat_run.font.size = Pt(12)
    cat_run.font.name = TNR
    catedra.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    docentes = doc.add_paragraph()
    doc_run = docentes.add_run("Docentes: Ing. Erica M. Milin, Ing. David Carlos Mammana")
    doc_run.bold = True
    doc_run.italic = True
    doc_run.font.size = Pt(12)
    doc_run.font.name = TNR
    docentes.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # --- ABSTRACT (Times New Roman, 10, negrita) ---
    p_abs_tit = doc.add_paragraph()
    tit_abs = p_abs_tit.add_run("Abstract")
    tit_abs.bold = True
    tit_abs.font.size = ABSTRACT_SIZE
    tit_abs.font.name = TNR
    
    abstract_text = (
        "Se presenta un estudio de simulación por eventos discretos de una plataforma SaaS que conecta "
        "técnicos con clientes. El objetivo es evaluar la viabilidad económica bajo distintas configuraciones "
        "de modelo de negocio (suscripción vs prepago), frecuencia de releases y presupuesto de marketing. "
        "Se ejecutó un benchmark de 5000 corridas por configuración sobre un horizonte de 10 años (3650 días), "
        "explorando 27 combinaciones de variables. Los resultados muestran que la frecuencia de implementaciones "
        "es el factor más determinante: releases semanales impiden alcanzar equilibrio en todas las configuraciones, "
        "mientras que releases mensuales o trimestrales permiten viabilidad con tasas de equilibrio entre 74% y 100%. "
        "La configuración más cost-effective (marketing 500 créditos/mes) que logra sostenibilidad es AB 50-50 con "
        "releases trimestrales, alcanzando beneficio medio de 241.970 créditos y equilibrio en 90% de corridas. "
        "La satisfacción de clientes suscriptores supera significativamente a la de prepago (96-97% vs 18-50%) "
        "en configuraciones con releases estables.")
    
    p_abs = doc.add_paragraph()
    abs_run = p_abs.add_run(abstract_text)
    abs_run.italic = True
    abs_run.font.size = ABSTRACT_SIZE
    abs_run.font.name = TNR
    p_abs.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # --- Palabras Clave ---
    p_kw_tit = doc.add_paragraph()
    kw_tit = p_kw_tit.add_run("Palabras Clave")
    kw_tit.bold = True
    kw_tit.font.size = ABSTRACT_SIZE
    kw_tit.font.name = TNR
    
    p_kw = doc.add_paragraph()
    kw_run = p_kw.add_run(
        "Simulación por eventos discretos, SaaS, plataforma técnica, modelo de negocio, "
        "suscripción, prepago, benchmark, viabilidad económica, satisfacción de clientes."
    )
    kw_run.font.size = ABSTRACT_SIZE
    kw_run.font.name = TNR
    
    # --- SECCIÓN 2 COLUMNAS (espaciado 1 cm entre columnas) ---
    _agregar_section_break_2_columnas(doc)
    
    doc.add_paragraph()
    
    # --- INTRODUCCIÓN ---
    doc.add_paragraph()
    p_int_tit = doc.add_paragraph()
    int_tit = p_int_tit.add_run("Introducción")
    int_tit.bold = True
    int_tit.font.size = SECCION_SIZE
    int_tit.font.name = TNR
    
    intro = (
        "Las plataformas que conectan técnicos con clientes representan un modelo de negocio SaaS en crecimiento. "
        "La viabilidad económica depende de múltiples factores: el modelo de monetización (suscripción vs prepago), "
        "la frecuencia de actualizaciones de software, el presupuesto de adquisición de clientes y la satisfacción "
        "que determina la retención. Evaluar estas variables en un entorno real implica costos y riesgos elevados; "
        "por ello, la simulación por eventos discretos permite explorar escenarios antes de implementarlos.\n\n"
        "El presente trabajo aborda un modelo de simulación día a día de una plataforma que ofrece tres modalidades "
        "de pago: suscripción mensual (10 créditos/mes con 15% descuento), prepago (paquete de minutos) y trabajo "
        "aislado. Los clientes llegan según un flujo de demanda con probabilidades de calendarización, arrepentimiento "
        "y falta. La satisfacción se calcula a partir de factores como conectividad, inestabilidad post-implementación "
        "y cumplimiento de citas. El objetivo es determinar qué configuraciones permiten alcanzar el punto de equilibrio "
        "y maximizar el beneficio neto en un horizonte de 10 años.")
    
    p_int = doc.add_paragraph()
    int_run = p_int.add_run(intro)
    int_run.font.size = TEXTO_SIZE
    int_run.font.name = TNR
    p_int.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # --- ELEMENTOS DEL TRABAJO Y METODOLOGÍA ---
    doc.add_paragraph()
    p_met_tit = doc.add_paragraph()
    met_tit = p_met_tit.add_run("Elementos del Trabajo y Metodología")
    met_tit.bold = True
    met_tit.font.size = SECCION_SIZE
    met_tit.font.name = TNR
    
    metodologia = (
        "Análisis teórico y clasificación de variables. Según la propuesta del trabajo, las variables se clasifican en:\n\n"
        "Datos: Clientes_Asiduos, Clientes_Nuevos, Tiempo_Entre_Llegadas, Tipo_De_Trabajo (Apps, IT, Dev), Duraciones, "
        "Nuevo_o_Preexistente, Tipo_Cliente_Nuevo, Se_Calendariza, Arrepentimiento_post_Calendarizacion, "
        "Falta_a_Reunion_Calendarizada, Insatisfaccion_base, Conectividad_pobre, Inestabilidad_por_Implementacion, "
        "entre otras.\n\n"
        "Control: N (frecuencia de implementaciones, impacto en estabilidad vs mejora continua), M (presupuesto de "
        "marketing, define cantidad diaria de clientes nuevos), AB_Testing (porcentaje 0-100% destinado a Suscripción, "
        "el resto a Prepago).\n\n"
        "Resultado: Beneficio_neto_acumulado, Beneficio_mensual_promedio, Beneficio_anualizado, T_Equilibrio "
        "(día en que beneficio > 0), Mejor_Trimestre (ventana de 120 días con mayor beneficio).\n\n"
        "Estado: T (reloj), clientes por modalidad (Suscripción, Prepago, Trabajo Aislado), Técnicos_Dev, Técnicos_AppsIT, "
        "Créditos_prepago_global, Ultimo_Dia_De_Implementacion, Contrataciones_pendientes.\n\n"
        "Metodología: Delta T / T = 1 día. Eventos propios del modelo: Contratación y Renuncia de técnicos, "
        "Llegada de clientes (nuevos y asiduos), Atender clientes, Reponer créditos, Implementación, "
        "Cobro de suscripciones, Pago de sueldos, Renovación de prepagos.\n\n"
        "Implementación. Se desarrolló la simulación en Python 3.8+. El modelo simula día a día el flujo de clientes "
        "(nuevos con presupuesto de marketing vs preexistentes), la asignación de trabajos a técnicos (TPLL/TPS), "
        "tipos de trabajo (Apps/IT, Desarrollo) con proporciones Dirichlet, y la calendarización con probabilidades "
        "de arrepentimiento (60%) y falta (5%). Los costos incluyen desarrollo de software, marketing mensual y salarios.\n\n"
        "Variables de experimento del benchmark: AB Testing (0-100, 100-0, 50-50), Frecuencia de releases "
        "(Semanales 7d, Mensuales 30d, Trimestrales 90d), Presupuesto de marketing (500, 1500, 2500 créditos/mes). "
        "Se ejecutaron 5000 corridas por cada una de las 27 configuraciones, horizonte T=3650 días. "
        "Adicionalmente, se ejecutó un benchmark de 3 casos relevantes con 1000 corridas y horizonte 5-6 años "
        "(cost-effective: 6 años para capturar equilibrio tardío): equilibrio más rápido (0-100/Trim/MKT2500), "
        "cost-effective (50-50/Trim/MKT500) y menos efectivo (100-0/Sem/MKT500). "
        "Métricas extraídas: beneficio final, día de equilibrio, tasa de equilibrio, satisfacción por tipo de cliente, "
        "entrada inicial (6 y 12 meses) y mejor trimestre.")
    
    p_met = doc.add_paragraph()
    met_run = p_met.add_run(metodologia)
    met_run.font.size = TEXTO_SIZE
    met_run.font.name = TNR
    p_met.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # --- RESULTADOS ---
    doc.add_paragraph()
    p_res_tit = doc.add_paragraph()
    res_tit = p_res_tit.add_run("Resultados")
    res_tit.bold = True
    res_tit.font.size = SECCION_SIZE
    res_tit.font.name = TNR
    
    resultados = (
        "La Tabla 1 resume los resultados principales. Diferencia entre mejor y peor escenario: 6,7 millones de créditos. "
        "Los releases semanales no alcanzan equilibrio en ninguna configuración (0% de corridas), excepto en dos casos "
        "extremos con marketing 2500 donde el equilibrio se alcanza en ~año 8-9. Releases mensuales y trimestrales "
        "alcanzan equilibrio en 74%-100% de corridas.\n\n"
        "Mejor configuración: 50-50 / Trimestrales / MKT 2500 con +5.116.487 créditos, 100% equilibrio, ~751 usuarios, "
        "83,8% satisfacción. Peor: 100-0 / Semanales / MKT 500 con -1.598.935 créditos, 0% equilibrio, 23 usuarios. "
        "Cost-effective (MKT 500): 50-50 / Trimestrales con ~242.000 créditos, 90% equilibrio, ~3,5 años al equilibrio.\n\n"
        "Ventanas de equilibrio (casos 5 años, 1000 corridas): Equilibrio más rápido (0-100/Trim/MKT2500): semana 49-75 "
        "(~1 a 1,4 años). Cost-effective (50-50/Trim/MKT500): semana 119-307 (~2,3 a 5,9 años). Menos efectivo "
        "(100-0/Sem/MKT500): no alcanzado en 5 años. La serie de beneficio acumulado (media ± 1σ) ilustra el cruce "
        "con cero en los casos viables y la trayectoria negativa persistente en el caso semanal.\n\n"
        "Caso extremo (anexo): Para evaluar el efecto de una inversión de marketing muy alta fuera del rango típico "
        "(500-4500 créditos/mes), se ejecutó un benchmark con M=10.000 créditos/mes, 50-50, releases trimestrales, "
        "5 años (1000 corridas). El beneficio final medio alcanza ~3,2 millones de créditos, 100% equilibrio, con punto "
        "de equilibrio muy temprano (semana 28-35, ~7 meses). La inversión masiva acelera la rentabilidad, aunque a un "
        "costo mensual 20 veces superior al caso cost-effective. La desviación estándar relativa es baja (~2,5%).\n\n"
        "Satisfacción: en releases semanales 69,2% general; en mensuales 82,7%; en trimestrales 83,9%. Suscriptores "
        "96-97% en configs estables; prepago 18-50% (más variable). El modelo híbrido 50-50 maximiza beneficio al "
        "combinar entrada inicial (prepago) con recurrencia (suscripción).")
    
    p_res = doc.add_paragraph()
    res_run = p_res.add_run(resultados)
    res_run.font.size = TEXTO_SIZE
    res_run.font.name = TNR
    p_res.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Tabla 1 (resumen)
    p_tabla = doc.add_paragraph()
    tab_tit = p_tabla.add_run("Tabla 1. Resumen de resultados por configuración (beneficio medio, % equilibrio, satisfacción)")
    tab_tit.italic = True
    tab_tit.font.size = Pt(10)
    tab_tit.font.name = TNR
    
    table = doc.add_table(rows=9, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Configuración"
    hdr[1].text = "Beneficio (créditos) / Eq% / Sat gen"
    for cell in hdr:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
                r.font.name = TNR
    
    datos_tabla = [
        ("0-100 / Semanales / MKT 500", "-1.578.818 / 0% / 67.7%"),
        ("0-100 / Mensuales / MKT 1500", "1.964.857 / 100% / 81.6%"),
        ("0-100 / Trimestrales / MKT 2500", "4.507.466 / 100% / 82.2%"),
        ("100-0 / Semanales / MKT 500", "-1.598.935 / 0% / 72.6%"),
        ("100-0 / Mensuales / MKT 2500", "5.063.660 / 100% / 83.7%"),
        ("50-50 / Trimestrales / MKT 500", "241.970 / 90% / 84.2% (cost-effective)"),
        ("50-50 / Trimestrales / MKT 2500", "5.116.487 / 100% / 83.8% (mejor)"),
        ("50-50 / Trimestrales / MKT 10.000", "~3.200.000 / 100% / (caso extremo, 5 años)"),
    ]
    for i, (cfg, val) in enumerate(datos_tabla, 1):
        row = table.rows[i]
        row.cells[0].text = cfg
        row.cells[1].text = val
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.name = TNR
    
    # Tabla 2 (casos 5 años - ventanas de equilibrio)
    p_tabla2 = doc.add_paragraph()
    tab2_tit = p_tabla2.add_run("Tabla 2. Ventanas de equilibrio de los casos relevantes y caso extremo (1000 corridas, 5-6 años)")
    tab2_tit.italic = True
    tab2_tit.font.size = Pt(10)
    tab2_tit.font.name = TNR
    
    table2 = doc.add_table(rows=5, cols=4)
    table2.style = "Table Grid"
    hdr2 = table2.rows[0].cells
    hdr2[0].text = "Caso"
    hdr2[1].text = "Configuración"
    hdr2[2].text = "Semana eq. (temprano-tardío)"
    hdr2[3].text = "Aprox. en años"
    for cell in hdr2:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
                r.font.name = TNR
    
    datos_tabla2 = [
        ("Equilibrio más rápido", "0-100 / Trim / MKT 2500", "49-75", "~1 a 1,4"),
        ("Cost-effective", "50-50 / Trim / MKT 500", "119-307", "~2,3 a 5,9"),
        ("Menos efectivo", "100-0 / Sem / MKT 500", "N/A", "No alcanzado"),
        ("Caso extremo (anexo)", "50-50 / Trim / MKT 10.000", "28-35", "~0,5 a 0,7"),
    ]
    for i, (caso, cfg, sem, anos) in enumerate(datos_tabla2, 1):
        row = table2.rows[i]
        row.cells[0].text = caso
        row.cells[1].text = cfg
        row.cells[2].text = sem
        row.cells[3].text = anos
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.name = TNR
    
    # --- PLACEHOLDER FIGURA 1 ---
    p_fig1 = doc.add_paragraph()
    fig1_run = p_fig1.add_run(
        "[PLACEHOLDER FIGURA 1 - INSERTAR IMAGEN: benchmark_10_anos/comparacion_AB.png]\n"
        "Figura 1. Beneficio final por tipo de AB Testing (Suscripción vs Prepago)."
    )
    fig1_run.italic = True
    fig1_run.font.size = Pt(10)
    fig1_run.font.name = TNR
    
    # --- PLACEHOLDER FIGURA 2 ---
    p_fig2 = doc.add_paragraph()
    fig2_run = p_fig2.add_run(
        "[PLACEHOLDER FIGURA 2 - INSERTAR IMAGEN: benchmark_10_anos/comparacion_Releases.png]\n"
        "Figura 2. Beneficio final por frecuencia de releases (factor crítico)."
    )
    fig2_run.italic = True
    fig2_run.font.size = Pt(10)
    fig2_run.font.name = TNR
    
    # --- PLACEHOLDER FIGURA 3 ---
    p_fig3 = doc.add_paragraph()
    fig3_run = p_fig3.add_run(
        "[PLACEHOLDER FIGURA 3 - INSERTAR IMAGEN: benchmark_10_anos/mejor_vs_peor_config.png]\n"
        "Figura 3. Comparación mejor vs peor configuración."
    )
    fig3_run.italic = True
    fig3_run.font.size = Pt(10)
    fig3_run.font.name = TNR
    
    # --- PLACEHOLDER FIGURA 4 ---
    p_fig4 = doc.add_paragraph()
    fig4_run = p_fig4.add_run(
        "[PLACEHOLDER FIGURA 4 - INSERTAR IMAGEN: benchmark_10_anos/equilibrio_por_config.png]\n"
        "Figura 4. Porcentaje de corridas que alcanzaron equilibrio por configuración."
    )
    fig4_run.italic = True
    fig4_run.font.size = Pt(10)
    fig4_run.font.name = TNR
    
    # --- PLACEHOLDER FIGURA 5 ---
    p_fig5 = doc.add_paragraph()
    fig5_run = p_fig5.add_run(
        "[PLACEHOLDER FIGURA 5 - INSERTAR IMAGEN: benchmark_10_anos/satisfaccion_por_release.png]\n"
        "Figura 5. Satisfacción de usuarios por frecuencia de releases."
    )
    fig5_run.italic = True
    fig5_run.font.size = Pt(10)
    fig5_run.font.name = TNR
    
    # --- PLACEHOLDER FIGURA 6 ---
    p_fig6 = doc.add_paragraph()
    fig6_run = p_fig6.add_run(
        "[PLACEHOLDER FIGURA 6 - INSERTAR IMAGEN: benchmark_10_anos/entrada_inicial_prepago_vs_suscripcion.png]\n"
        "Figura 6. Ingresos primeros 6 meses: Prepago vs Suscripción por AB Testing."
    )
    fig6_run.italic = True
    fig6_run.font.size = Pt(10)
    fig6_run.font.name = TNR
    
    # --- PLACEHOLDERS CASOS 5 AÑOS ---
    p_fig7 = doc.add_paragraph()
    fig7_run = p_fig7.add_run(
        "[PLACEHOLDER FIGURA 7 - INSERTAR IMAGEN: benchmark_10_anos/casos_5_anos/equilibrio_antes/serie_beneficio_acumulado.png]\n"
        "Figura 7. Beneficio acumulado medio ± 1σ. Equilibrio más rápido (0-100/Trim/MKT2500). Cruce con cero: semana 49-75."
    )
    fig7_run.italic = True
    fig7_run.font.size = Pt(10)
    fig7_run.font.name = TNR
    
    p_fig8 = doc.add_paragraph()
    fig8_run = p_fig8.add_run(
        "[PLACEHOLDER FIGURA 8 - INSERTAR IMAGEN: benchmark_10_anos/casos_5_anos/cost_effective/serie_beneficio_acumulado.png]\n"
        "Figura 8. Beneficio acumulado medio ± 1σ. Cost-effective (50-50/Trim/MKT500). Cruce con cero: semana 119-307."
    )
    fig8_run.italic = True
    fig8_run.font.size = Pt(10)
    fig8_run.font.name = TNR
    
    p_fig9 = doc.add_paragraph()
    fig9_run = p_fig9.add_run(
        "[PLACEHOLDER FIGURA 9 - INSERTAR IMAGEN: benchmark_10_anos/casos_5_anos/menos_efectivo/serie_beneficio_acumulado.png]\n"
        "Figura 9. Beneficio acumulado medio ± 1σ. Menos efectivo (100-0/Sem/MKT500). Equilibrio no alcanzado en 5 años."
    )
    fig9_run.italic = True
    fig9_run.font.size = Pt(10)
    fig9_run.font.name = TNR
    
    p_fig10 = doc.add_paragraph()
    fig10_run = p_fig10.add_run(
        "[PLACEHOLDER FIGURA 10 - INSERTAR IMAGEN: benchmark_10_anos/casos_5_anos/menos_efectivo/ejemplo/perdidas_clientes.png]\n"
        "Figura 10. Pérdidas de clientes por razón (corrida representativa). Caso menos efectivo: motivos del fracaso."
    )
    fig10_run.italic = True
    fig10_run.font.size = Pt(10)
    fig10_run.font.name = TNR
    
    p_fig11 = doc.add_paragraph()
    fig11_run = p_fig11.add_run(
        "[PLACEHOLDER FIGURA 11 - INSERTAR IMAGEN: benchmark_10_anos/casos_5_anos/cost_effective/ejemplo/satisfaccion_general.png]\n"
        "Figura 11. Evolución de satisfacción general (corrida representativa). Caso cost-effective viable."
    )
    fig11_run.italic = True
    fig11_run.font.size = Pt(10)
    fig11_run.font.name = TNR
    
    # --- PLACEHOLDER FIGURA 12 (caso extremo) ---
    p_fig12 = doc.add_paragraph()
    fig12_run = p_fig12.add_run(
        "[PLACEHOLDER FIGURA 12 - INSERTAR IMAGEN: graficos_benchmark/caso_extremo_mkt10000/serie_beneficio_acumulado.png]\n"
        "Figura 12. Beneficio acumulado medio ± 1σ. Caso extremo (50-50/Trim/MKT 10.000). Equilibrio ~7 meses (semana 28-35). 1000 corridas, 5 años."
    )
    fig12_run.italic = True
    fig12_run.font.size = Pt(10)
    fig12_run.font.name = TNR
    
    doc.add_paragraph()
    
    # --- DISCUSIÓN ---
    doc.add_paragraph()
    p_dis_tit = doc.add_paragraph()
    dis_tit = p_dis_tit.add_run("Discusión")
    dis_tit.bold = True
    dis_tit.font.size = SECCION_SIZE
    dis_tit.font.name = TNR
    
    discusion = (
        "El hallazgo crítico (Informe de Hallazgos): la frecuencia de releases es el factor más determinante. "
        "Releases semanales reducen la satisfacción en ~14 puntos porcentuales y destruyen valor en TODAS las "
        "configuraciones. Cambiar de semanal a trimestral aporta +3,7M de beneficio promedio; cambiar de 100-0 "
        "a 50-50 solo +0,6M. Mensaje clave: la estabilidad de la plataforma es más importante que el modelo de negocio.\n\n"
        "La satisfacción de suscriptores es muy estable (96-97%) en configs mensuales/trimestrales. El prepago "
        "no \"salva\" configuraciones semanales: aunque 0-100 genera más ingresos iniciales, con releases semanales "
        "nunca alcanza equilibrio. El modelo híbrido 50-50 equilibra entrada inicial (prepago) y recurrencia (suscripción).\n\n"
        "Config cost-effective (MKT 500, releases trimestrales): viabilidad con baja inversión. Los costos fijos "
        "se compensan con ingresos recurrentes sin flujo masivo de clientes. Las series de beneficio acumulado "
        "de los 3 casos relevantes (Fig. 7-9) ilustran el contraste: equilibrio en ~1 año (0-100/Trim/MKT2500), "
        "en ~2,3-5,9 años (50-50/Trim/MKT500) o nunca (100-0/Sem/MKT500). El análisis de pérdidas de clientes "
        "(Fig. 10) en el caso menos efectivo revela las razones del fracaso (arrepentimiento, calendarización, etc.). "
        "Configuraciones a evitar: cualquier combinación con releases semanales (0% probabilidad de éxito).\n\n"
        "El caso extremo (MKT 10.000, Fig. 12) confirma que el presupuesto de marketing acelera el punto de equilibrio: "
        "con el doble del máximo del rango estándar, el equilibrio se alcanza en ~7 meses frente a 2,3-5,9 años del "
        "cost-effective. Sin embargo, el retorno marginal decrece: pasar de 500 a 10.000 créditos/mes multiplica el "
        "beneficio final por ~13, pero la inversión mensual se multiplica por 20. Sugiere un punto óptimo intermedio. "
        "El caso extremo sirve como cota superior teórica, no como estrategia recomendada por su alto costo.")
    
    p_dis = doc.add_paragraph()
    dis_run = p_dis.add_run(discusion)
    dis_run.font.size = TEXTO_SIZE
    dis_run.font.name = TNR
    p_dis.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # --- CONCLUSIÓN ---
    doc.add_paragraph()
    p_con_tit = doc.add_paragraph()
    con_tit = p_con_tit.add_run("Conclusión")
    con_tit.bold = True
    con_tit.font.size = SECCION_SIZE
    con_tit.font.name = TNR
    
    conclusion = (
        "La simulación por eventos discretos permitió evaluar la viabilidad económica de una plataforma SaaS "
        "de servicios técnicos bajo 27 configuraciones. Los resultados indican que: (1) la frecuencia de releases "
        "es el factor más determinante—releases semanales impiden el equilibrio en todas las configuraciones; "
        "(2) releases mensuales o trimestrales permiten viabilidad con tasas de equilibrio entre 74% y 100%; "
        "(3) la configuración más cost-effective (marketing 500 créditos/mes) es AB 50-50 con releases trimestrales; "
        "(4) la satisfacción de suscriptores supera significativamente a la de prepago en configuraciones estables; "
        "(5) el caso extremo con marketing 10.000 créditos/mes acelera el equilibrio (~7 meses) y genera beneficios "
        "elevados (~3,2 M créditos), pero con retorno marginal decreciente—se recomienda priorizar marketing moderado "
        "(500-2500 créditos/mes). Se recomienda priorizar releases mensuales o trimestrales y un modelo de monetización "
        "mixto para maximizar tanto la entrada inicial como la retención.")
    
    p_con = doc.add_paragraph()
    con_run = p_con.add_run(conclusion)
    con_run.font.size = TEXTO_SIZE
    con_run.font.name = TNR
    p_con.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # --- REFERENCIAS ---
    doc.add_paragraph()
    p_ref_tit = doc.add_paragraph()
    ref_tit = p_ref_tit.add_run("Referencias")
    ref_tit.bold = True
    ref_tit.font.size = REF_SIZE
    ref_tit.font.name = TNR
    
    refs = [
        "[1] Day, R. A. (1991). How to Write and Publish a Scientific Paper. Cambridge University Press.",
        "[2] American National Standards Institute. Definition of abstract.",
        "[3] Documentación del proyecto: https://github.com/apacay/Simu-vAri/wiki",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        r = p.add_run(ref)
        r.font.size = REF_SIZE
        r.font.name = TNR
    
    # --- AGRADECIMIENTOS ---
    doc.add_paragraph()
    p_agr_tit = doc.add_paragraph()
    agr_tit = p_agr_tit.add_run("Agradecimientos")
    agr_tit.bold = True
    agr_tit.font.size = REF_SIZE
    agr_tit.font.name = TNR
    
    p_agr = doc.add_paragraph()
    agr_run = p_agr.add_run(
        "A los docentes de la Cátedra Simulación, Ing. Erica M. Milin e Ing. David Carlos Mammana, "
        "por la guía y supervisión del trabajo."
    )
    agr_run.font.size = REF_SIZE
    agr_run.font.name = TNR
    
    # --- DATOS DE CONTACTO (Times New Roman 10, negrita título; contenido en cursiva) ---
    doc.add_paragraph()
    p_contact_tit = doc.add_paragraph()
    contact_tit = p_contact_tit.add_run("Datos de Contacto")
    contact_tit.bold = True
    contact_tit.font.size = REF_SIZE
    contact_tit.font.name = TNR
    
    p_contact = doc.add_paragraph()
    contact_run = p_contact.add_run(
        "Claudio Hernán Luciano Sonntag (120580-8), Ariel Gastón Pacay (121237-0), Alison Michella Reynoso (156899-1). "
        "Universidad Tecnológica Nacional, Facultad Regional Buenos Aires. Medrano 951, C1179AAQ CABA."
    )
    contact_run.italic = True
    contact_run.font.size = REF_SIZE
    contact_run.font.name = TNR
    
    # Guardar
    out_path = Path(__file__).parent / "Paper_Simulacion_Plataforma_SaaS.docx"
    doc.save(str(out_path))
    print(f"Paper generado: {out_path}")
    return out_path


if __name__ == "__main__":
    crear_paper()
